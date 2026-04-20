from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.services.config import ROOT_DIR

COVER_LETTER_OUTPUT_DIR = ROOT_DIR / "artifacts" / "cover_letters"
DEFAULT_FILE_STEM = "cover_letter"
DESIGN_SETTINGS_PATH = ROOT_DIR / "config" / "cover_letter_design.json"
SOFFICE_CANDIDATES = [
    shutil.which("soffice"),
    shutil.which("libreoffice"),
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]


def _load_design_settings() -> dict[str, Any]:
    defaults: dict[str, Any] = {
        "font_name": "Times New Roman",
        "body_font_size_pt": 11.0,
        "name_font_size_pt": 13.0,
        "contact_font_size_pt": 10.0,
        "margins_inches": {"top": 0.85, "bottom": 0.85, "left": 0.9, "right": 0.9},
        "spacing": {
            "header_after_pt": 10.0,
            "date_after_pt": 8.0,
            "recipient_after_pt": 8.0,
            "salutation_after_pt": 8.0,
            "body_after_pt": 8.0,
            "signature_before_pt": 16.0,
        },
        "line_spacing": {"body": 1.0833333333},
        "salutation": "Dear Hiring Team,",
        "closing": "Sincerely,",
    }
    if DESIGN_SETTINGS_PATH.exists():
        loaded = json.loads(DESIGN_SETTINGS_PATH.read_text(encoding="utf-8"))
        if isinstance(loaded, dict):
            merged = defaults | loaded
            merged["margins_inches"] = defaults["margins_inches"] | loaded.get("margins_inches", {})
            merged["spacing"] = defaults["spacing"] | loaded.get("spacing", {})
            merged["line_spacing"] = defaults["line_spacing"] | loaded.get("line_spacing", {})
            return merged
    return defaults


DESIGN = _load_design_settings()


def _safe_stem(value: str | None) -> str:
    candidate = (value or DEFAULT_FILE_STEM).strip()
    candidate = Path(candidate).stem
    candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", candidate).strip("._-")
    return candidate or DEFAULT_FILE_STEM


def _normalize_text(value: Any) -> str:
    return " ".join(str(value or "").strip().split())


def _contact_line(contact: dict[str, Any]) -> str:
    values = []
    for key in ["location", "phone", "email", "linkedin"]:
        text = _normalize_text(contact.get(key, ""))
        text = re.sub(r"^https?://", "", text)
        text = text.replace("www.", "")
        if text:
            values.append(text)
    return " | ".join(values)


def _recipient_lines(company_context: dict[str, Any], explicit_lines: list[str] | None) -> list[str]:
    if explicit_lines:
        return [_normalize_text(line) for line in explicit_lines if _normalize_text(line)]

    lines = [
        _normalize_text(company_context.get("company", "")),
        _normalize_text(
            company_context.get("team")
            or company_context.get("department")
            or company_context.get("business_unit")
            or company_context.get("recipient_unit", "")
        ),
    ]
    return [line for line in lines if line]


def _body_paragraphs(cover_letter: str) -> list[str]:
    parts = re.split(r"\n\s*\n", str(cover_letter or "").strip())
    cleaned = [_normalize_text(part) for part in parts if _normalize_text(part)]
    return cleaned


def _set_run_font(run, size_pt: float, bold: bool = False) -> None:
    run.font.name = str(DESIGN["font_name"])
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_paragraph(paragraph, *, after_pt: float = 0, before_pt: float = 0, line_spacing: float | None = None, align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after_pt)
    fmt.space_before = Pt(before_pt)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing


def _add_paragraph(document: Document, text: str, *, size_pt: float, bold: bool = False, align=None, after_pt: float = 0, before_pt: float = 0, line_spacing: float | None = None) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph(paragraph, after_pt=after_pt, before_pt=before_pt, line_spacing=line_spacing, align=align)
    run = paragraph.add_run(text)
    _set_run_font(run, size_pt, bold=bold)


def _write_docx(
    docx_path: Path,
    *,
    name: str,
    contact: dict[str, Any],
    date_text: str,
    recipient_lines: list[str],
    salutation: str,
    body_paragraphs: list[str],
    closing: str,
    signature_name: str,
) -> None:
    document = Document()
    section = document.sections[0]
    margins = DESIGN["margins_inches"]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(float(margins["top"]))
    section.bottom_margin = Inches(float(margins["bottom"]))
    section.left_margin = Inches(float(margins["left"]))
    section.right_margin = Inches(float(margins["right"]))

    normal = document.styles["Normal"]
    normal.font.name = str(DESIGN["font_name"])
    normal.font.size = Pt(float(DESIGN["body_font_size_pt"]))

    spacing = DESIGN["spacing"]
    _add_paragraph(
        document,
        _normalize_text(name).upper(),
        size_pt=float(DESIGN["name_font_size_pt"]),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    contact_line = _contact_line(contact)
    if contact_line:
        _add_paragraph(
            document,
            contact_line,
            size_pt=float(DESIGN["contact_font_size_pt"]),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after_pt=float(spacing["header_after_pt"]),
        )

    _add_paragraph(document, date_text, size_pt=float(DESIGN["body_font_size_pt"]), after_pt=float(spacing["date_after_pt"]))
    for index, line in enumerate(recipient_lines):
        _add_paragraph(
            document,
            line,
            size_pt=float(DESIGN["body_font_size_pt"]),
            after_pt=float(spacing["recipient_after_pt"]) if index == len(recipient_lines) - 1 else 0,
        )

    _add_paragraph(document, salutation, size_pt=float(DESIGN["body_font_size_pt"]), after_pt=float(spacing["salutation_after_pt"]))
    for paragraph_text in body_paragraphs:
        _add_paragraph(
            document,
            paragraph_text,
            size_pt=float(DESIGN["body_font_size_pt"]),
            after_pt=float(spacing["body_after_pt"]),
            line_spacing=float(DESIGN["line_spacing"]["body"]),
        )

    _add_paragraph(document, closing, size_pt=float(DESIGN["body_font_size_pt"]))
    _add_paragraph(
        document,
        signature_name,
        size_pt=float(DESIGN["body_font_size_pt"]),
        before_pt=float(spacing["signature_before_pt"]),
    )
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def _find_soffice() -> str | None:
    for candidate in SOFFICE_CANDIDATES:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> str:
    soffice = _find_soffice()
    if soffice:
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        expected_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"
        if completed.returncode == 0 and expected_pdf.exists():
            if expected_pdf != pdf_path:
                shutil.move(str(expected_pdf), str(pdf_path))
            return "soffice"
    return "reportlab_fallback"


def _write_fallback_pdf(
    pdf_path: Path,
    *,
    name: str,
    contact_line: str = "",
    date_text: str = "",
    recipient_lines: list[str] | None = None,
    salutation: str = "",
    body_paragraphs: list[str] | None = None,
    closing: str = "",
    signature_name: str = "",
) -> None:
    recipient_lines = recipient_lines or []
    body_paragraphs = body_paragraphs or []
    buffer = BytesIO()
    margins = DESIGN["margins_inches"]
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=float(margins["left"]) * inch,
        rightMargin=float(margins["right"]) * inch,
        topMargin=float(margins["top"]) * inch,
        bottomMargin=float(margins["bottom"]) * inch,
    )
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        name="CoverHeader",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=float(DESIGN["name_font_size_pt"]),
        alignment=1,
        spaceAfter=0,
    )
    contact_style = ParagraphStyle(
        name="CoverContact",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=float(DESIGN["contact_font_size_pt"]),
        alignment=1,
        spaceAfter=float(DESIGN["spacing"]["header_after_pt"]),
    )
    body_style = ParagraphStyle(
        name="CoverBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=float(DESIGN["body_font_size_pt"]),
        leading=float(DESIGN["body_font_size_pt"]) * float(DESIGN["line_spacing"]["body"]),
        spaceAfter=float(DESIGN["spacing"]["body_after_pt"]),
    )
    story: list[Any] = [
        Paragraph(_normalize_text(name).upper(), header_style),
    ]
    if contact_line:
        story.append(Paragraph(contact_line, contact_style))
    story.append(Paragraph(date_text, body_style))
    for index, line in enumerate(recipient_lines):
        story.append(Paragraph(line, body_style))
        if index == len(recipient_lines) - 1:
            story.append(Spacer(1, float(DESIGN["spacing"]["recipient_after_pt"])))
    story.append(Paragraph(salutation, body_style))
    for paragraph_text in body_paragraphs:
        story.append(Paragraph(paragraph_text, body_style))
    story.append(Paragraph(closing, body_style))
    story.append(Spacer(1, float(DESIGN["spacing"]["signature_before_pt"])))
    story.append(Paragraph(signature_name, body_style))
    doc.build(story)
    pdf_path.write_bytes(buffer.getvalue())


def _page_count_from_reader(pdf_path: Path) -> int:
    return len(PdfReader(str(pdf_path)).pages)


def build_cover_letter_document(
    *,
    cover_letter: str,
    contact: dict[str, Any],
    company_context: dict[str, Any] | None = None,
    date_text: str = "",
    recipient_lines: list[str] | None = None,
    signature_name: str = "",
    output_dir: str | Path | None = None,
    output_filename: str | None = None,
) -> dict[str, Any]:
    company_context = company_context or {}
    final_output_dir = Path(output_dir) if output_dir else COVER_LETTER_OUTPUT_DIR
    final_output_dir.mkdir(parents=True, exist_ok=True)
    file_stem = _safe_stem(output_filename)
    final_docx_path = final_output_dir / f"{file_stem}.docx"
    final_pdf_path = final_output_dir / f"{file_stem}.pdf"

    name = str(company_context.get("applicant_name", signature_name)).strip() or signature_name
    letter_date = _normalize_text(date_text) or datetime.now().strftime("%B %d, %Y").replace(" 0", " ")
    lines = _recipient_lines(company_context, recipient_lines)
    salutation = _normalize_text(company_context.get("salutation", DESIGN["salutation"])) or str(DESIGN["salutation"])
    closing = _normalize_text(company_context.get("closing", DESIGN["closing"])) or str(DESIGN["closing"])
    paragraphs = _body_paragraphs(cover_letter)
    contact_line = _contact_line(contact)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        docx_path = temp_path / f"{file_stem}.docx"
        pdf_path = temp_path / f"{file_stem}.pdf"
        _write_docx(
            docx_path,
            name=name,
            contact=contact,
            date_text=letter_date,
            recipient_lines=lines,
            salutation=salutation,
            body_paragraphs=paragraphs,
            closing=closing,
            signature_name=signature_name,
        )
        conversion_method = _convert_docx_to_pdf(docx_path, pdf_path)
        if conversion_method == "reportlab_fallback":
            _write_fallback_pdf(
                pdf_path,
                name=name,
                contact_line=contact_line,
                date_text=letter_date,
                recipient_lines=lines,
                salutation=salutation,
                body_paragraphs=paragraphs,
                closing=closing,
                signature_name=signature_name,
            )
        page_count = _page_count_from_reader(pdf_path)
        shutil.copy2(docx_path, final_docx_path)
        shutil.copy2(pdf_path, final_pdf_path)

    return {
        "docx_path": str(final_docx_path),
        "pdf_path": str(final_pdf_path),
        "page_count": page_count,
        "fits_on_one_page": page_count <= 1,
        "conversion_method": conversion_method,
        "page_count_method": "pypdf",
        "applied_design": DESIGN,
    }
