from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from backend.services.config import ROOT_DIR

ATS_OUTPUT_DIR = ROOT_DIR / "artifacts" / "resumes"
DEFAULT_FILE_STEM = "resume"
DESIGN_SETTINGS_PATH = ROOT_DIR / "config" / "resume_design.json"


def _load_design_settings() -> dict[str, Any]:
    defaults: dict[str, Any] = {
        "font_name": "Times New Roman",
        "content_indent_inch": 0.08,
        "layout_attempts": [
            {"font_size_pt": 10.5, "max_bullets_per_role": 4, "margin_inches": 0.40, "max_experience_items": 5, "max_projects": 2, "max_publications": 1, "max_leadership_items": 2},
            {"font_size_pt": 10.5, "max_bullets_per_role": 3, "margin_inches": 0.40, "max_experience_items": 5, "max_projects": 2, "max_publications": 1, "max_leadership_items": 2},
            {"font_size_pt": 10.5, "max_bullets_per_role": 3, "margin_inches": 0.40, "max_experience_items": 5, "max_projects": 1, "max_publications": 1, "max_leadership_items": 1},
            {"font_size_pt": 10.5, "max_bullets_per_role": 3, "margin_inches": 0.40, "max_experience_items": 4, "max_projects": 1, "max_publications": 1, "max_leadership_items": 1},
            {"font_size_pt": 10.25, "max_bullets_per_role": 3, "margin_inches": 0.40, "max_experience_items": 4, "max_projects": 1, "max_publications": 1, "max_leadership_items": 1},
            {"font_size_pt": 10.0, "max_bullets_per_role": 3, "margin_inches": 0.38, "max_experience_items": 4, "max_projects": 1, "max_publications": 0, "max_leadership_items": 1},
        ],
        "publication_display": {
            "show_status": False,
            "strip_locations": ["Philadelphia"],
        },
    }
    if DESIGN_SETTINGS_PATH.exists():
        loaded = json.loads(DESIGN_SETTINGS_PATH.read_text(encoding="utf-8"))
        if isinstance(loaded, dict):
            publication_display = defaults["publication_display"] | loaded.get("publication_display", {})
            return defaults | loaded | {"publication_display": publication_display}
    return defaults


DESIGN_SETTINGS = _load_design_settings()
FONT_NAME = str(DESIGN_SETTINGS["font_name"])
CONTENT_INDENT_INCH = float(DESIGN_SETTINGS["content_indent_inch"])
CONTENT_INDENT_PT = CONTENT_INDENT_INCH * 72


@dataclass(frozen=True)
class LayoutAttempt:
    font_size_pt: float
    max_bullets_per_role: int
    margin_inches: float
    max_experience_items: int
    max_projects: int
    max_publications: int
    max_leadership_items: int


LAYOUT_ATTEMPTS = [LayoutAttempt(**attempt) for attempt in DESIGN_SETTINGS["layout_attempts"]]

SOFFICE_CANDIDATES = [
    shutil.which("soffice"),
    shutil.which("libreoffice"),
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]


def _styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ResumeSection",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=colors.black,
            spaceBefore=4,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=11,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeHeader",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=17,
            spaceAfter=6,
        )
    )
    return styles


def render_pdf(title: str, sections: list[dict]) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )
    styles = _styles()
    story = [Paragraph(title, styles["ResumeHeader"]), Spacer(1, 4)]

    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        if heading:
            story.append(Paragraph(heading, styles["ResumeSection"]))
        if isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"- {item}", styles["ResumeBody"]))
        else:
            story.append(Paragraph(str(content), styles["ResumeBody"]))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buffer.getvalue()


def _safe_stem(value: str | None) -> str:
    candidate = (value or DEFAULT_FILE_STEM).strip()
    candidate = Path(candidate).stem
    candidate = re.sub(r"[^A-Za-z0-9._-]+", "_", candidate).strip("._-")
    return candidate or DEFAULT_FILE_STEM


def _normalize_bullet_text(bullet: Any) -> str:
    if isinstance(bullet, dict):
        return str(bullet.get("text", "")).strip()
    return str(bullet).strip()


def _trim_resume_for_attempt(tailored_resume: dict[str, Any], attempt: LayoutAttempt) -> dict[str, Any]:
    trimmed: dict[str, Any] = {
        key: value for key, value in tailored_resume.items()
    }
    role_bucket = str(tailored_resume.get("role_bucket", "")).strip()
    priority_weights = tailored_resume.get("priority_engine", {}).get("weights", {})
    experience_items = tailored_resume.get("experience", [])[: attempt.max_experience_items]
    trimmed["experience"] = []
    for exp in experience_items:
        exp_copy = dict(exp)
        bullets = [_normalize_bullet_text(bullet) for bullet in exp.get("bullets", [])]
        exp_copy["bullets"] = [bullet for bullet in bullets[: attempt.max_bullets_per_role] if bullet]
        trimmed["experience"].append(exp_copy)
    trimmed["projects"] = trimmed.get("projects", [])[: attempt.max_projects]
    if attempt.max_publications > 0:
        if int(priority_weights.get("publications", 0)) >= 7:
            trimmed["publications"] = trimmed.get("publications", [])[: attempt.max_publications]
        else:
            trimmed["publications"] = trimmed.get("publications", [])[:1]
    else:
        trimmed["publications"] = []
    trimmed["leadership"] = trimmed.get("leadership", [])[: attempt.max_leadership_items]
    skills_display = trimmed.get("skills_display", [])
    if isinstance(skills_display, list):
        compact_skills = []
        for index, group in enumerate(skills_display):
            group_copy = dict(group)
            item_limit = 6 if index == 0 else 5
            if attempt.font_size_pt <= 10.25:
                item_limit -= 1
            if attempt.max_bullets_per_role <= 2:
                item_limit -= 1 if index > 0 else 0
            group_copy["items"] = group_copy.get("items", [])[: max(item_limit, 4)]
            compact_skills.append(group_copy)
        trimmed["skills_display"] = compact_skills
    return trimmed


def _set_run_font(run, font_name: str, font_size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold


def _set_paragraph_format(paragraph, *, before: float = 0, after: float = 0, line: float = 1.0) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def _add_text_line(
    document: Document,
    text: str,
    *,
    font_size_pt: float,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
    left_indent: float = 0,
) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    _set_paragraph_format(paragraph, before=before, after=after, line=line)
    if left_indent:
        paragraph.paragraph_format.left_indent = Inches(left_indent)
    run = paragraph.add_run(text)
    _set_run_font(run, FONT_NAME, font_size_pt, bold=bold)


def _add_section_heading(document: Document, heading: str, font_size_pt: float) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_format(paragraph, before=3.5, after=1.5, line=1.0)
    paragraph.paragraph_format.left_indent = Inches(CONTENT_INDENT_INCH)
    run = paragraph.add_run(heading.upper())
    _set_run_font(run, FONT_NAME, font_size_pt, bold=True)
    paragraph_borders = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    border.append(bottom)
    paragraph_borders.append(border)


def _build_contact_line(contact: dict[str, Any]) -> str:
    def _display_value(value: Any) -> str:
        text = str(value).strip()
        text = re.sub(r"^https?://", "", text)
        text = text.replace("www.", "")
        return text

    values = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("linkedin", ""),
        contact.get("github", ""),
        contact.get("location", ""),
    ]
    return " | ".join(_display_value(value) for value in values if str(value).strip())


def _display_section_heading(tailored_resume: dict[str, Any], key: str, fallback: str) -> str:
    headings = tailored_resume.get("section_headings", {})
    if isinstance(headings, dict):
        return str(headings.get(key, fallback))
    return fallback


def _add_left_right_line(
    document: Document,
    left_text: str,
    right_text: str,
    *,
    font_size_pt: float,
    bold: bool = False,
    after: float = 0,
    margin_inches: float,
) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_format(paragraph, before=0, after=after, line=1.0)
    paragraph.paragraph_format.left_indent = Inches(CONTENT_INDENT_INCH)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(8.5 - (margin_inches * 2) - CONTENT_INDENT_INCH),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
    )
    run = paragraph.add_run(f"{left_text}\t{right_text}".rstrip())
    _set_run_font(run, FONT_NAME, font_size_pt, bold=bold)


def _add_bullet_line(document: Document, text: str, font_size_pt: float) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_format(paragraph, before=0, after=0, line=0.98)
    paragraph.paragraph_format.left_indent = Inches(CONTENT_INDENT_INCH)
    run = paragraph.add_run(f"• {text}")
    _set_run_font(run, FONT_NAME, font_size_pt)


def _add_section_end_space(document: Document, amount: float = 2.0) -> None:
    spacer = document.add_paragraph()
    _set_paragraph_format(spacer, before=0, after=amount, line=1.0)


def _display_dates(text: str) -> str:
    value = str(text or "").strip()
    value = value.replace("â€“", "–").replace("â€”", "—")
    value = value.replace(" - ", " – ")
    value = re.sub(r"\s*-\s*", " – ", value)
    return value


def _display_publication_venue(publication: dict[str, Any]) -> str:
    venue = str(publication.get("venue", "")).strip()
    strip_locations = DESIGN_SETTINGS.get("publication_display", {}).get("strip_locations", [])
    for location in strip_locations:
        venue = venue.replace(f", {location}", "")
        venue = venue.replace(location, "")
    venue = venue.strip(" ,|")
    return venue


def _write_docx_resume(docx_path: Path, tailored_resume: dict[str, Any], attempt: LayoutAttempt) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(attempt.margin_inches)
    section.bottom_margin = Inches(attempt.margin_inches)
    section.left_margin = Inches(attempt.margin_inches)
    section.right_margin = Inches(attempt.margin_inches)

    normal_style = document.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(attempt.font_size_pt)

    name = tailored_resume.get("name", "[user]")
    _add_text_line(
        document,
        str(name).upper(),
        font_size_pt=max(attempt.font_size_pt + 3.5, 13),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0.25,
    )

    contact_line = _build_contact_line(tailored_resume.get("contact", {}))
    if contact_line:
        _add_text_line(
            document,
            contact_line,
            font_size_pt=max(attempt.font_size_pt - 0.75, 8.75),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after=1.75,
        )

    education = tailored_resume.get("education", [])
    if education:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "education", "Education"),
            attempt.font_size_pt,
        )
        for edu in education:
            _add_left_right_line(
                document,
                str(edu.get("school", "")).strip(),
                str(edu.get("location", tailored_resume.get("contact", {}).get("location", ""))).strip(),
                font_size_pt=attempt.font_size_pt,
                bold=True,
                margin_inches=attempt.margin_inches,
            )
            _add_left_right_line(
                document,
                " | ".join(
                    part
                    for part in [
                        str(edu.get("program", edu.get("degree", ""))).strip(),
                        f"GPA: {edu.get('gpa')}" if edu.get("gpa") else "",
                    ]
                    if part
                ),
                _display_dates(str(edu.get("graduation", "")).strip()),
                font_size_pt=attempt.font_size_pt,
                margin_inches=attempt.margin_inches,
            )
            awards = [str(award).strip() for award in edu.get("awards", []) if str(award).strip()]
            if awards:
                paragraph = document.add_paragraph()
                _set_paragraph_format(paragraph, before=0, after=0.4, line=1.0)
                paragraph.paragraph_format.left_indent = Inches(CONTENT_INDENT_INCH)
                awards_label = paragraph.add_run("Awards: ")
                _set_run_font(awards_label, FONT_NAME, max(attempt.font_size_pt - 0.35, 8.5), bold=True)
                awards_text = paragraph.add_run("; ".join(awards))
                _set_run_font(awards_text, FONT_NAME, max(attempt.font_size_pt - 0.35, 8.5))
        _add_section_end_space(document, 0.4)

    experience = tailored_resume.get("experience", [])
    if experience:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "experience", "PROFESSIONAL EXPERIENCE"),
            attempt.font_size_pt,
        )
        for exp in experience:
            _add_left_right_line(
                document,
                str(exp.get("company", "")).strip(),
                str(exp.get("location", "")).strip(),
                font_size_pt=attempt.font_size_pt,
                bold=True,
                margin_inches=attempt.margin_inches,
            )
            _add_left_right_line(
                document,
                str(exp.get("title", "")).strip(),
                _display_dates(str(exp.get("dates", "")).strip()),
                font_size_pt=attempt.font_size_pt,
                margin_inches=attempt.margin_inches,
                after=0.25,
            )
            for bullet in exp.get("bullets", []):
                bullet_text = _normalize_bullet_text(bullet)
                if bullet_text:
                    _add_bullet_line(document, bullet_text, attempt.font_size_pt)
            _add_section_end_space(document, 1.65)

    projects = tailored_resume.get("projects", [])
    if projects:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "projects", "Project"),
            attempt.font_size_pt,
        )
        for project in projects:
            name = str(project.get("name", "")).strip()
            summary = str(project.get("display_summary", project.get("summary", project.get("description", "")))).strip()
            line = f"{name}: {summary}" if summary else name
            _add_text_line(document, line, font_size_pt=max(attempt.font_size_pt - 0.15, 8.75), after=0)
        _add_section_end_space(document, 0.9)

    publications = tailored_resume.get("publications", [])
    if publications:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "publications", "PUBLICATIONS"),
            attempt.font_size_pt,
        )
        for publication in publications:
            title_line = str(publication.get("title", "")).strip()
            venue_line = _display_publication_venue(publication)
            _add_left_right_line(
                document,
                title_line,
                venue_line,
                font_size_pt=max(attempt.font_size_pt - 0.1, 8.9),
                bold=True,
                margin_inches=attempt.margin_inches,
            )
            summary = str(publication.get("display_summary", publication.get("summary", ""))).strip()
            if summary:
                _add_text_line(document, summary, font_size_pt=max(attempt.font_size_pt - 0.2, 8.7))
        _add_section_end_space(document, 0.9)

    skills_display = tailored_resume.get("skills_display", [])
    skills = tailored_resume.get("skills", {})
    if skills_display or skills:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "skills", "Skills"),
            attempt.font_size_pt,
        )
        if skills_display:
            for group in skills_display:
                label = str(group.get("label", "")).strip()
                values = [str(item).strip() for item in group.get("items", []) if str(item).strip()]
                if values:
                    _add_text_line(
                        document,
                        f"{label}: {', '.join(values)}",
                        font_size_pt=max(attempt.font_size_pt - 0.25, 8.75),
                    )
        elif isinstance(skills, dict):
            for category, items in skills.items():
                values = [str(item).strip() for item in items if str(item).strip()]
                if values:
                    _add_text_line(
                        document,
                        f"{str(category).replace('_', ' ').title()}: {', '.join(values)}",
                        font_size_pt=max(attempt.font_size_pt - 0.25, 8.75),
                    )
        _add_section_end_space(document, 0.35)

    leadership = tailored_resume.get("leadership", [])
    if leadership:
        _add_section_heading(
            document,
            _display_section_heading(tailored_resume, "leadership", "Leadership"),
            attempt.font_size_pt,
        )
        for item in leadership:
            header_left = " | ".join(
                part for part in [str(item.get("role", "")).strip(), str(item.get("organization", "")).strip()] if part
            )
            header_right = " | ".join(
                part for part in [_display_dates(str(item.get("dates", "")).strip()), str(item.get("location", "")).strip()] if part
            )
            _add_left_right_line(
                document,
                header_left,
                header_right,
                font_size_pt=attempt.font_size_pt,
                bold=True,
                margin_inches=attempt.margin_inches,
            )
            if item.get("description"):
                _add_bullet_line(document, str(item.get("description", "")).strip(), max(attempt.font_size_pt - 0.15, 8.75))
            _add_section_end_space(document, 0.8)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def _find_soffice() -> str | None:
    for candidate in SOFFICE_CANDIDATES:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path, tailored_resume: dict[str, Any], attempt: LayoutAttempt) -> str:
    soffice = _find_soffice()
    if soffice:
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        expected_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"
        if completed.returncode == 0 and expected_pdf.exists():
            if expected_pdf != pdf_path:
                shutil.move(str(expected_pdf), str(pdf_path))
            return "soffice"

    _write_fallback_pdf(pdf_path, tailored_resume, attempt)
    return "reportlab_fallback"


def _pdftk_page_count(pdf_path: Path) -> int | None:
    pdftk = shutil.which("pdftk")
    if not pdftk:
        return None
    completed = subprocess.run(
        [pdftk, str(pdf_path), "dump_data"],
        capture_output=True,
        text=True,
        check=False,
    )
    if completed.returncode != 0:
        return None
    match = re.search(r"NumberOfPages:\s+(\d+)", completed.stdout)
    return int(match.group(1)) if match else None


def _page_count_from_reader(pdf_path: Path) -> int:
    reader = PdfReader(str(pdf_path))
    return len(reader.pages)


def _get_page_count(pdf_path: Path) -> tuple[int, str]:
    pdftk_count = _pdftk_page_count(pdf_path)
    if pdftk_count is not None:
        return pdftk_count, "pdftk"
    return _page_count_from_reader(pdf_path), "pypdf"


def _write_fallback_pdf(pdf_path: Path, tailored_resume: dict[str, Any], attempt: LayoutAttempt) -> None:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=attempt.margin_inches * inch,
        rightMargin=attempt.margin_inches * inch,
        topMargin=attempt.margin_inches * inch,
        bottomMargin=attempt.margin_inches * inch,
    )
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        name="AtsHeader",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=max(attempt.font_size_pt + 3.5, 13),
        alignment=1,
        spaceAfter=1,
        leading=max(attempt.font_size_pt + 4, 14),
    )
    contact_style = ParagraphStyle(
        name="AtsContact",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=max(attempt.font_size_pt - 0.75, 8.75),
        alignment=1,
        spaceAfter=3,
        leading=max(attempt.font_size_pt + 0.5, 10),
    )
    section_style = ParagraphStyle(
        name="AtsSection",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=attempt.font_size_pt,
        leftIndent=CONTENT_INDENT_PT,
        spaceBefore=3.5,
        spaceAfter=1.5,
        leading=max(attempt.font_size_pt + 0.5, 10),
    )
    body_style = ParagraphStyle(
        name="AtsBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=attempt.font_size_pt,
        leftIndent=CONTENT_INDENT_PT,
        spaceAfter=0.75,
        leading=max(attempt.font_size_pt + 0.75, 10.5),
    )
    small_style = ParagraphStyle(
        name="AtsSmall",
        parent=body_style,
        fontSize=max(attempt.font_size_pt - 0.25, 8.6),
        leading=max(attempt.font_size_pt + 0.5, 10),
    )
    bullet_style = ParagraphStyle(
        name="AtsBullet",
        parent=body_style,
        spaceAfter=0,
        leading=max(attempt.font_size_pt + 0.35, 9.75),
    )

    story: list[Any] = []
    story.append(Paragraph(str(tailored_resume.get("name", "[user]")).upper(), header_style))
    contact_line = _build_contact_line(tailored_resume.get("contact", {}))
    if contact_line:
        story.append(Paragraph(contact_line, contact_style))

    def section_header(text: str) -> list[Any]:
        return [
            Paragraph(text, section_style),
            HRFlowable(width=doc.width - CONTENT_INDENT_PT, thickness=0.8, color=colors.HexColor("#666666"), spaceBefore=0, spaceAfter=1.5, hAlign="RIGHT"),
        ]

    def two_col(left: str, right: str, *, bold: bool = False) -> Table:
        left_text = f"<b>{left}</b>" if bold and left else left
        right_text = f"<b>{right}</b>" if bold and right else right
        right_style = ParagraphStyle(
            name=f"AtsRight_{'bold' if bold else 'base'}",
            parent=small_style if not bold else body_style,
            alignment=2,
            leading=max(attempt.font_size_pt + 0.5, 10),
        )
        table = Table(
            [[Paragraph(left_text, body_style), Paragraph(right_text, right_style)]],
            colWidths=[doc.width * 0.78, doc.width * 0.22],
        )
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (0, -1), CONTENT_INDENT_PT),
                    ("LEFTPADDING", (1, 0), (1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
                ]
            )
        )
        return table

    education = tailored_resume.get("education", [])
    if education:
        story.extend(section_header(_display_section_heading(tailored_resume, "education", "Education").upper()))
        for edu in education:
            story.append(two_col(str(edu.get("school", "")).strip(), str(edu.get("location", tailored_resume.get("contact", {}).get("location", ""))).strip(), bold=True))
            story.append(
                two_col(
                    " | ".join(
                        part
                        for part in [
                            str(edu.get("program", edu.get("degree", ""))).strip(),
                            f"GPA: {edu.get('gpa')}" if edu.get("gpa") else "",
                        ]
                        if part
                    ),
                    _display_dates(str(edu.get("graduation", "")).strip()),
                )
            )
            awards = [str(award).strip() for award in edu.get("awards", []) if str(award).strip()]
            if awards:
                story.append(Paragraph(f"<b>Awards:</b> {'; '.join(awards)}", small_style))
        story.append(Spacer(1, 0.75))

    experience = tailored_resume.get("experience", [])
    if experience:
        story.extend(section_header(_display_section_heading(tailored_resume, "experience", "Experience").upper()))
        for exp in experience:
            story.append(two_col(str(exp.get("company", "")).strip(), str(exp.get("location", "")).strip(), bold=True))
            story.append(two_col(str(exp.get("title", "")).strip(), _display_dates(str(exp.get("dates", "")).strip())))
            for bullet in exp.get("bullets", []):
                bullet_text = _normalize_bullet_text(bullet)
                if bullet_text:
                    story.append(Paragraph(f"• {bullet_text}", bullet_style))
            story.append(Spacer(1, 3.0))

    projects = tailored_resume.get("projects", [])
    if projects:
        story.extend(section_header(_display_section_heading(tailored_resume, "projects", "Project").upper()))
        for project in projects:
            name = str(project.get("name", "")).strip()
            summary = str(project.get("display_summary", project.get("summary", project.get("description", "")))).strip()
            story.append(Paragraph(f"<b>{name}</b>: {summary}" if summary else f"<b>{name}</b>", small_style))
        story.append(Spacer(1, 2))

    publications = tailored_resume.get("publications", [])
    if publications:
        story.extend(section_header(_display_section_heading(tailored_resume, "publications", "Research Papers").upper()))
        for publication in publications:
            venue_line = _display_publication_venue(publication)
            story.append(two_col(str(publication.get("title", "")).strip(), venue_line, bold=True))
            summary = str(publication.get("display_summary", publication.get("summary", ""))).strip()
            if summary:
                story.append(Paragraph(summary, small_style))
        story.append(Spacer(1, 2))

    skills_display = tailored_resume.get("skills_display", [])
    skills = tailored_resume.get("skills", {})
    if skills_display or skills:
        story.extend(section_header(_display_section_heading(tailored_resume, "skills", "Skills").upper()))
        if skills_display:
            for group in skills_display:
                label = str(group.get("label", "")).strip()
                values = [str(item).strip() for item in group.get("items", []) if str(item).strip()]
                if values:
                    story.append(Paragraph(f"<b>{label}:</b> {', '.join(values)}", small_style))
        elif isinstance(skills, dict):
            for category, items in skills.items():
                values = [str(item).strip() for item in items if str(item).strip()]
                if values:
                    story.append(Paragraph(f"<b>{str(category).replace('_', ' ').title()}:</b> {', '.join(values)}", small_style))

    leadership = tailored_resume.get("leadership", [])
    if leadership:
        story.extend(section_header(_display_section_heading(tailored_resume, "leadership", "Leadership").upper()))
        for item in leadership:
            story.append(
                two_col(
                    " | ".join(
                        part for part in [str(item.get("role", "")).strip(), str(item.get("organization", "")).strip()] if part
                    ),
                    " | ".join(
                        part for part in [_display_dates(str(item.get("dates", "")).strip()), str(item.get("location", "")).strip()] if part
                    ),
                    bold=True,
                )
            )
            if item.get("description"):
                story.append(Paragraph(f"• {str(item.get('description', '')).strip()}", bullet_style))
            story.append(Spacer(1, 1.5))

    doc.build(story)
    pdf_path.write_bytes(buffer.getvalue())


def _default_layout_steps(attempt: LayoutAttempt) -> list[str]:
    steps: list[str] = []
    if attempt.font_size_pt < 10.5:
        steps.append(f"font_reduced_to_{attempt.font_size_pt}pt")
    if attempt.max_bullets_per_role < 4:
        steps.append(f"bullet_limit_reduced_to_{attempt.max_bullets_per_role}")
    if attempt.margin_inches < 0.40:
        steps.append(f"margins_reduced_to_{attempt.margin_inches:.2f}in")
    if attempt.max_experience_items < 4:
        steps.append(f"experience_limit_reduced_to_{attempt.max_experience_items}")
    if attempt.max_leadership_items < 1:
        steps.append(f"leadership_limit_reduced_to_{attempt.max_leadership_items}")
    return steps


def build_ats_resume_pdf(
    tailored_resume: dict[str, Any],
    output_dir: str | Path | None = None,
    output_filename: str | None = None,
) -> dict[str, Any]:
    final_output_dir = Path(output_dir) if output_dir else ATS_OUTPUT_DIR
    final_output_dir.mkdir(parents=True, exist_ok=True)
    file_stem = _safe_stem(output_filename)
    final_docx_path = final_output_dir / f"{file_stem}.docx"
    final_pdf_path = final_output_dir / f"{file_stem}.pdf"

    best_candidate: dict[str, Any] | None = None
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        for index, attempt in enumerate(LAYOUT_ATTEMPTS, start=1):
            working_resume = _trim_resume_for_attempt(tailored_resume, attempt)
            docx_path = temp_path / f"{file_stem}_{index}.docx"
            pdf_path = temp_path / f"{file_stem}_{index}.pdf"

            _write_docx_resume(docx_path, working_resume, attempt)
            conversion_method = _convert_docx_to_pdf(docx_path, pdf_path, working_resume, attempt)
            page_count, page_count_method = _get_page_count(pdf_path)

            candidate = {
                "attempt": attempt,
                "docx_path": docx_path,
                "pdf_path": pdf_path,
                "page_count": page_count,
                "page_count_method": page_count_method,
                "conversion_method": conversion_method,
            }
            if best_candidate is None or candidate["page_count"] < best_candidate["page_count"]:
                best_candidate = candidate

            if page_count <= 1:
                best_candidate = candidate
                break
        assert best_candidate is not None
        shutil.copy2(best_candidate["docx_path"], final_docx_path)
        shutil.copy2(best_candidate["pdf_path"], final_pdf_path)

        attempt = best_candidate["attempt"]
        return {
            "docx_path": str(final_docx_path),
            "pdf_path": str(final_pdf_path),
            "page_count": best_candidate["page_count"],
            "fits_on_one_page": best_candidate["page_count"] <= 1,
            "applied_layout": {
                "font_size_pt": attempt.font_size_pt,
                "max_bullets_per_role": attempt.max_bullets_per_role,
                "margin_inches": attempt.margin_inches,
                "max_experience_items": attempt.max_experience_items,
                "max_projects": attempt.max_projects,
                "max_publications": attempt.max_publications,
                "max_leadership_items": attempt.max_leadership_items,
                "overflow_steps_applied": _default_layout_steps(attempt),
            },
            "conversion_method": best_candidate["conversion_method"],
            "page_count_method": best_candidate["page_count_method"],
        }
